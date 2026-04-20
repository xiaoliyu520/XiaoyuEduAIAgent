from fastapi import APIRouter, Depends, UploadFile, File, Query
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, or_
import json
import hashlib
from datetime import datetime
from typing import Optional
import io

from app.core.database import get_db
from app.models.database import User, KnowledgeBase, KnowledgeDocument, KnowledgeGap, DocumentVersion
from app.models.schemas import (
    KnowledgeBaseCreate, KnowledgeBaseResponse, KnowledgeGapResponse,
    KnowledgeGapResolve, KnowledgeGapIgnore, ResponseBase, DocumentResponse, DocumentVersionResponse,
    PageResponse,
)
from app.api.deps import get_current_user, require_admin
from app.mcp.milvus.client import (
    insert_documents, delete_collection, get_collection_stats, ensure_collection,
    delete_documents_by_metadata, search, update_document_status,
)
from app.mcp.milvus.bm25 import get_bm25_index
from app.core.minio import upload_file

router = APIRouter(prefix="/knowledge", tags=["知识库"])


def extract_text_from_file(file_data: bytes, file_type: str) -> str:
    if file_type == "pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_data))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text.strip())
            return "\n\n".join(text_parts)
        except Exception as e:
            return ""
    elif file_type == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_data))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(text_parts)
        except Exception as e:
            return ""
    else:
        return file_data.decode("utf-8", errors="ignore")


@router.post("/bases", response_model=ResponseBase)
async def create_knowledge_base(
    data: KnowledgeBaseCreate,
    current_user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    collection_name = f"kb_{hashlib.md5(f'{data.tenant_id}_{data.name}'.encode()).hexdigest()[:12]}"
    ensure_collection(collection_name)

    kb = KnowledgeBase(
        name=data.name,
        description=data.description,
        collection_name=collection_name,
        tenant_id=data.tenant_id,
    )
    db.add(kb)
    await db.commit()
    await db.refresh(kb)
    return ResponseBase(data=KnowledgeBaseResponse.model_validate(kb).model_dump())


@router.get("/bases", response_model=ResponseBase)
async def list_knowledge_bases(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeBase))
    kbs = result.scalars().all()
    items = [KnowledgeBaseResponse.model_validate(kb).model_dump() for kb in kbs]
    return ResponseBase(data=items)


@router.delete("/bases/{kb_id}", response_model=ResponseBase)
async def delete_knowledge_base(
    kb_id: int,
    current_user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == kb_id))
    kb = result.scalar_one_or_none()
    if not kb:
        raise ValueError("知识库不存在")
    delete_collection(kb.collection_name)
    doc_result = await db.execute(
        select(KnowledgeDocument).where(KnowledgeDocument.kb_id == kb_id)
    )
    for doc in doc_result.scalars().all():
        await db.delete(doc)
    await db.delete(kb)
    await db.commit()
    return ResponseBase(message="知识库已删除")


@router.post("/bases/{kb_id}/upload", response_model=ResponseBase)
async def upload_document(
    kb_id: int,
    file: UploadFile = File(...),
    current_user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == kb_id))
    kb = result.scalar_one_or_none()
    if not kb:
        raise ValueError("知识库不存在")

    file_data = await file.read()
    content_hash = hashlib.sha256(file_data).hexdigest()
    file_size = len(file_data)

    existing_doc = await db.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.kb_id == kb_id,
            KnowledgeDocument.title == file.filename,
        )
    )
    existing_doc = existing_doc.scalar_one_or_none()

    if existing_doc:
        if existing_doc.content_hash == content_hash:
            return ResponseBase(message="文档内容未变化，无需更新", data={"unchanged": True})

        return await _update_document(
            db=db,
            kb=kb,
            existing_doc=existing_doc,
            file_data=file_data,
            file_name=file.filename,
            content_type=file.content_type or "application/octet-stream",
            content_hash=content_hash,
            file_size=file_size,
        )

    object_name = f"knowledge/{kb.collection_name}/{file.filename}"
    await upload_file(object_name, file_data, file.content_type or "application/octet-stream")

    file_ext = file.filename.split(".")[-1].lower() if "." in file.filename else "txt"
    content = extract_text_from_file(file_data, file_ext)
    
    if not content.strip():
        return ResponseBase(message="文档内容为空或无法解析", data={"empty": True})
    
    chunks = _split_text(content, chunk_size=500, overlap=50)

    if chunks:
        chunk_ids = [f"{file.filename}_chunk_{i}" for i in range(len(chunks))]
        await insert_documents(
            collection_name=kb.collection_name,
            contents=chunks,
            metadatas=[{"source": file.filename, "chunk_index": i, "doc_title": file.filename, "doc_status": "active"} for i in range(len(chunks))],
            chunk_ids=chunk_ids,
        )

        bm25 = get_bm25_index()
        bm25_docs = [{"content": c, "chunk_id": chunk_ids[i], "metadata": {"source": file.filename, "chunk_index": i, "doc_title": file.filename, "doc_status": "active"}} for i, c in enumerate(chunks)]
        await bm25.add_documents(kb.collection_name, bm25_docs)

    doc = KnowledgeDocument(
        kb_id=kb_id,
        title=file.filename,
        file_path=object_name,
        file_type=file.filename.split(".")[-1] if "." in file.filename else "txt",
        file_size=file_size,
        chunk_count=len(chunks),
        content_hash=content_hash,
        version=1,
    )
    db.add(doc)
    await db.flush()

    version_record = DocumentVersion(
        doc_id=doc.id,
        version=1,
        file_path=object_name,
        content_hash=content_hash,
        chunk_count=len(chunks),
        change_type="created",
        change_summary="文档首次上传",
    )
    db.add(version_record)

    kb.doc_count += len(chunks)
    await db.commit()

    return ResponseBase(data={
        "chunks": len(chunks),
        "file": file.filename,
        "version": 1,
        "is_new": True,
    })


async def _update_document(
    db: AsyncSession,
    kb: KnowledgeBase,
    existing_doc: KnowledgeDocument,
    file_data: bytes,
    file_name: str,
    content_type: str,
    content_hash: str,
    file_size: int,
) -> ResponseBase:
    old_content_hash = existing_doc.content_hash
    old_chunk_count = existing_doc.chunk_count

    object_name = f"knowledge/{kb.collection_name}/{file_name}_v{existing_doc.version + 1}"
    await upload_file(object_name, file_data, content_type)

    file_ext = file_name.split(".")[-1].lower() if "." in file_name else "txt"
    content = extract_text_from_file(file_data, file_ext)
    
    if not content.strip():
        return ResponseBase(message="文档内容为空或无法解析", data={"empty": True})
    
    new_chunks = _split_text(content, chunk_size=500, overlap=50)

    await delete_documents_by_metadata(kb.collection_name, {"source": file_name})

    bm25 = get_bm25_index()
    await bm25.remove_document(kb.collection_name, file_name)

    if new_chunks:
        doc_status = existing_doc.status if existing_doc.status else "active"
        chunk_ids = [f"{file_name}_chunk_{i}" for i in range(len(new_chunks))]
        await insert_documents(
            collection_name=kb.collection_name,
            contents=new_chunks,
            metadatas=[{"source": file_name, "chunk_index": i, "doc_title": file_name, "doc_status": doc_status} for i in range(len(new_chunks))],
            chunk_ids=chunk_ids,
        )

        bm25_docs = [{"content": c, "chunk_id": chunk_ids[i], "metadata": {"source": file_name, "chunk_index": i, "doc_title": file_name, "doc_status": doc_status}} for i, c in enumerate(new_chunks)]
        await bm25.add_documents(kb.collection_name, bm25_docs)

    chunk_diff = len(new_chunks) - old_chunk_count
    if chunk_diff > 0:
        change_summary = f"新增 {chunk_diff} 个文本块"
    elif chunk_diff < 0:
        change_summary = f"减少 {abs(chunk_diff)} 个文本块"
    else:
        change_summary = "内容已更新，文本块数量不变"

    existing_doc.file_path = object_name
    existing_doc.file_size = file_size
    existing_doc.chunk_count = len(new_chunks)
    existing_doc.content_hash = content_hash
    existing_doc.version += 1
    existing_doc.updated_at = datetime.utcnow()

    version_record = DocumentVersion(
        doc_id=existing_doc.id,
        version=existing_doc.version,
        file_path=object_name,
        content_hash=content_hash,
        chunk_count=len(new_chunks),
        change_type="updated",
        change_summary=change_summary,
    )
    db.add(version_record)

    kb.doc_count = kb.doc_count - old_chunk_count + len(new_chunks)
    await db.commit()

    return ResponseBase(data={
        "chunks": len(new_chunks),
        "file": file_name,
        "version": existing_doc.version,
        "is_new": False,
        "change_summary": change_summary,
    })


@router.get("/bases/{kb_id}/documents", response_model=ResponseBase)
async def list_documents(
    kb_id: int,
    keyword: Optional[str] = Query(None, description="搜索关键词"),
    status: Optional[str] = Query(None, description="文档状态"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == kb_id))
    kb = result.scalar_one_or_none()
    if not kb:
        raise ValueError("知识库不存在")

    query = select(KnowledgeDocument).where(KnowledgeDocument.kb_id == kb_id)

    if keyword:
        query = query.where(KnowledgeDocument.title.ilike(f"%{keyword}%"))
    if status:
        query = query.where(KnowledgeDocument.status == status)

    count_query = select(func.count()).select_from(query.subquery())
    total_result = await db.execute(count_query)
    total = total_result.scalar()

    query = query.order_by(KnowledgeDocument.updated_at.desc())
    query = query.offset((page - 1) * page_size).limit(page_size)

    result = await db.execute(query)
    documents = result.scalars().all()

    items = [DocumentResponse.model_validate(doc).model_dump() for doc in documents]

    return ResponseBase(data=PageResponse(
        total=total,
        page=page,
        page_size=page_size,
        items=items,
    ).model_dump())


@router.get("/documents/{doc_id}", response_model=ResponseBase)
async def get_document(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise ValueError("文档不存在")
    return ResponseBase(data=DocumentResponse.model_validate(doc).model_dump())


@router.get("/documents/{doc_id}/versions", response_model=ResponseBase)
async def list_document_versions(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise ValueError("文档不存在")

    versions_result = await db.execute(
        select(DocumentVersion)
        .where(DocumentVersion.doc_id == doc_id)
        .order_by(DocumentVersion.version.desc())
    )
    versions = versions_result.scalars().all()

    items = [DocumentVersionResponse.model_validate(v).model_dump() for v in versions]
    return ResponseBase(data=items)


@router.delete("/documents/{doc_id}", response_model=ResponseBase)
async def delete_document(
    doc_id: int,
    current_user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise ValueError("文档不存在")

    kb_result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == doc.kb_id))
    kb = kb_result.scalar_one_or_none()

    if kb:
        await delete_documents_by_metadata(kb.collection_name, {"source": doc.title})
        bm25 = get_bm25_index()
        await bm25.remove_document(kb.collection_name, doc.title)
        kb.doc_count -= doc.chunk_count
        if kb.doc_count < 0:
            kb.doc_count = 0

    versions_result = await db.execute(
        select(DocumentVersion).where(DocumentVersion.doc_id == doc_id)
    )
    for version in versions_result.scalars().all():
        await db.delete(version)

    await db.delete(doc)
    await db.commit()

    return ResponseBase(message="文档已删除")


@router.put("/documents/{doc_id}/status", response_model=ResponseBase)
async def update_document_status_api(
    doc_id: int,
    status: str = Query(..., description="文档状态: active/archived"),
    current_user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise ValueError("文档不存在")

    doc.status = status
    await db.commit()

    kb_result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == doc.kb_id))
    kb = kb_result.scalar_one_or_none()
    if kb:
        await update_document_status(kb.collection_name, doc.title, status)
        bm25 = get_bm25_index()
        await bm25.update_document_status(kb.collection_name, doc.title, status)

    return ResponseBase(data=DocumentResponse.model_validate(doc).model_dump())


@router.get("/gaps", response_model=ResponseBase)
async def list_knowledge_gaps(
    status: str = "open",
    kb_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    query = select(KnowledgeGap).where(KnowledgeGap.status == status)
    if kb_id:
        query = query.where(KnowledgeGap.kb_id == kb_id)
    query = query.order_by(KnowledgeGap.created_at.desc())
    result = await db.execute(query)
    gaps = result.scalars().all()
    
    kb_map = {}
    kb_ids = set(g.kb_id for g in gaps if g.kb_id)
    if kb_ids:
        kb_result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id.in_(kb_ids)))
        for kb in kb_result.scalars().all():
            kb_map[kb.id] = kb.name
    
    items = []
    for g in gaps:
        item = KnowledgeGapResponse.model_validate(g).model_dump()
        item["kb_name"] = kb_map.get(g.kb_id)
        items.append(item)
    
    return ResponseBase(data=items)


@router.put("/gaps/{gap_id}/resolve", response_model=ResponseBase)
async def resolve_knowledge_gap(
    gap_id: int,
    data: KnowledgeGapResolve,
    current_user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    import logging
    logger = logging.getLogger(__name__)
    
    result = await db.execute(select(KnowledgeGap).where(KnowledgeGap.id == gap_id))
    gap = result.scalar_one_or_none()
    if not gap:
        raise ValueError("知识缺口不存在")
    
    if not gap.kb_id:
        gap.answer = data.answer
        gap.status = "resolved"
        gap.resolved_at = datetime.now()
        await db.commit()
        logger.warning(f"知识缺口 {gap_id} 未关联知识库，仅更新状态")
        return ResponseBase(
            message="知识缺口未关联知识库，答案已保存但未补录到知识库",
            data=KnowledgeGapResponse.model_validate(gap).model_dump()
        )
    
    kb_result = await db.execute(select(KnowledgeBase).where(KnowledgeBase.id == gap.kb_id))
    kb = kb_result.scalar_one_or_none()
    if not kb:
        raise ValueError("关联的知识库不存在")
    
    doc_title = "知识补录"
    new_entry = f"\n\n---\n\n## 问题：{gap.question}\n\n**答案：**{data.answer}"
    
    existing_doc_result = await db.execute(
        select(KnowledgeDocument).where(
            KnowledgeDocument.kb_id == gap.kb_id,
            KnowledgeDocument.title == doc_title
        )
    )
    existing_doc = existing_doc_result.scalar_one_or_none()
    
    bm25 = get_bm25_index()
    
    if existing_doc:
        logger.info(f"追加到现有补录文档 '{doc_title}'")
        
        await delete_documents_by_metadata(
            kb.collection_name,
            {"doc_title": doc_title}
        )
        await bm25.remove_document(kb.collection_name, doc_title)
        
        existing_content = ""
        old_gap_results = await db.execute(
            select(KnowledgeGap).where(
                KnowledgeGap.kb_id == gap.kb_id,
                KnowledgeGap.status == "resolved",
                KnowledgeGap.id != gap_id
            ).order_by(KnowledgeGap.resolved_at)
        )
        old_gaps = old_gap_results.scalars().all()
        for old_gap in old_gaps:
            if old_gap.answer and not old_gap.answer.startswith("[忽略原因]"):
                existing_content += f"\n\n---\n\n## 问题：{old_gap.question}\n\n**答案：**{old_gap.answer}"
        
        full_content = existing_content + new_entry
        content_hash = hashlib.sha256(full_content.encode()).hexdigest()
        
        chunks = _split_text(full_content, chunk_size=500, overlap=50)
        if chunks:
            chunk_ids = [f"{doc_title}_chunk_{i}" for i in range(len(chunks))]
            await insert_documents(
                collection_name=kb.collection_name,
                contents=chunks,
                metadatas=[{
                    "source": doc_title,
                    "chunk_index": i,
                    "doc_title": doc_title,
                    "doc_status": "active",
                } for i in range(len(chunks))],
                chunk_ids=chunk_ids,
            )
            
            bm25_docs = [{
                "content": c,
                "chunk_id": chunk_ids[i],
                "metadata": {
                    "source": doc_title,
                    "chunk_index": i,
                    "doc_title": doc_title,
                    "doc_status": "active"
                }
            } for i, c in enumerate(chunks)]
            await bm25.add_documents(kb.collection_name, bm25_docs)
        
        existing_doc.content_hash = content_hash
        existing_doc.chunk_count = len(chunks)
        existing_doc.file_size = len(full_content.encode())
        existing_doc.version += 1
        existing_doc.updated_at = datetime.now()
        
        version_record = DocumentVersion(
            doc_id=existing_doc.id,
            version=existing_doc.version,
            file_path=f"gap://collection",
            content_hash=content_hash,
            chunk_count=len(chunks),
            change_type="updated",
            change_summary=f"追加知识缺口: {gap.question[:30]}",
        )
        db.add(version_record)
        
        logger.info(f"文档 '{doc_title}' 追加成功，共 {len(chunks)} 个分块")
    else:
        logger.info(f"创建新补录文档 '{doc_title}'")
        
        full_content = f"# 知识补录{new_entry}"
        content_hash = hashlib.sha256(full_content.encode()).hexdigest()
        
        chunks = _split_text(full_content, chunk_size=500, overlap=50)
        if chunks:
            chunk_ids = [f"{doc_title}_chunk_{i}" for i in range(len(chunks))]
            await insert_documents(
                collection_name=kb.collection_name,
                contents=chunks,
                metadatas=[{
                    "source": doc_title,
                    "chunk_index": i,
                    "doc_title": doc_title,
                    "doc_status": "active",
                } for i in range(len(chunks))],
                chunk_ids=chunk_ids,
            )
            
            bm25_docs = [{
                "content": c,
                "chunk_id": chunk_ids[i],
                "metadata": {
                    "source": doc_title,
                    "chunk_index": i,
                    "doc_title": doc_title,
                    "doc_status": "active"
                }
            } for i, c in enumerate(chunks)]
            await bm25.add_documents(kb.collection_name, bm25_docs)
        
        doc = KnowledgeDocument(
            kb_id=gap.kb_id,
            title=doc_title,
            file_path="gap://collection",
            file_type="txt",
            file_size=len(full_content.encode()),
            chunk_count=len(chunks),
            content_hash=content_hash,
            version=1,
        )
        db.add(doc)
        await db.flush()
        
        version_record = DocumentVersion(
            doc_id=doc.id,
            version=1,
            file_path="gap://collection",
            content_hash=content_hash,
            chunk_count=len(chunks),
            change_type="created",
            change_summary=f"知识补录文档创建: {gap.question[:30]}",
        )
        db.add(version_record)
        
        kb.doc_count += len(chunks)
        
        logger.info(f"文档 '{doc_title}' 创建成功，共 {len(chunks)} 个分块")
    
    gap.answer = data.answer
    gap.status = "resolved"
    gap.resolved_at = datetime.now()
    await db.commit()
    
    return ResponseBase(
        message=f"已补录到知识库 '{kb.name}'",
        data=KnowledgeGapResponse.model_validate(gap).model_dump()
    )


@router.put("/gaps/{gap_id}/ignore", response_model=ResponseBase)
async def ignore_knowledge_gap(
    gap_id: int,
    data: KnowledgeGapIgnore = None,
    current_user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(select(KnowledgeGap).where(KnowledgeGap.id == gap_id))
    gap = result.scalar_one_or_none()
    if not gap:
        raise ValueError("知识缺口不存在")
    gap.status = "ignored"
    if data and data.reason:
        gap.answer = f"[忽略原因] {data.reason}"
    gap.resolved_at = datetime.now()
    await db.commit()
    return ResponseBase(data=KnowledgeGapResponse.model_validate(gap).model_dump())


def _split_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start = end - overlap
    return chunks
